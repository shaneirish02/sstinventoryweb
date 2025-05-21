from django.shortcuts import render, get_object_or_404, redirect
from .models import Item, Category, Sale, UserProfile, Supplier, CartItem, StockTransaction, SaleItem
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login
from django.db.models import Sum
from django.contrib.auth.models import User
from .forms import UserForm, ItemForm
from django.http import JsonResponse
import datetime
import json
from django.contrib import messages
from .models import UserProfile
from rest_framework.authtoken.models import Token

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.utils.text import slugify
from django.http import HttpResponse
from django.utils.dateparse import parse_date
from datetime import datetime, timedelta

from django.core.exceptions import ValidationError
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import logout
from django.views.decorators.csrf import csrf_exempt
from .forms import SupplierForm
from django.template.loader import render_to_string
from django.db.models import Q
from django.utils import timezone
from xhtml2pdf import pisa
from docx import Document
import io
import os
from django.conf import settings
from reportlab.lib import colors
from django.utils.timezone import now
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.contrib.auth import authenticate
from .serializer import UserSerializer, ItemSerializer, StockTransactionSerializer
from datetime import datetime, timedelta
from .models import StockTransaction



def logout_view(request):
    logout(request)
    return redirect('user_login')

@csrf_exempt
def checkout(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)

            payment_received = float(data.get('payment_received', 0))
            cart = request.session.get('cart', [])

            if not cart:
                return JsonResponse({'success': False, 'message': 'Cart is empty.'})

            total_amount = sum(item['price'] * item['quantity'] for item in cart)

            if payment_received < total_amount:
                return JsonResponse({'success': False, 'message': 'Insufficient funds.'})

            # Create Sale record
            sale = Sale.objects.create(
                total_price=total_amount,
                date=timezone.now().date()
            )

            print("Cart in session:", cart)

            # Process each item in the cart
            for cart_item in cart:
                print("Cart item:", cart_item)
                item_id = cart_item.get('id')
                if not item_id:
                    return JsonResponse({'success': False, 'message': 'Missing item ID in cart.', 'cart_item': cart_item})

                try:
                    item = Item.objects.get(id=item_id)
                except Item.DoesNotExist:
                    return JsonResponse({'success': False, 'message': f'Item with ID {item_id} not found.'})

                # Create sale item entry
                sale.items.create(
                    item=item,
                    quantity=cart_item['quantity'],
                    price=item.price,
                    subtotal=item.price * cart_item['quantity']
                )

                # Create stock transaction entry for each item
                StockTransaction.objects.create(
                    item=item,
                    transaction_type='OUT',  # 'OUT' means item is being checked out
                    quantity=cart_item['quantity'],
                    unit=item.unit,
                    date=timezone.now(),
                    remarks='Checked out via POS'
                )

                # Deduct stock
                item.quantity -= cart_item['quantity']
                item.save()

            # Clear cart after successful checkout
            request.session['cart'] = []
            request.session.modified = True

            change = payment_received - total_amount

            return JsonResponse({
                'success': True,
                'message': 'Transaction successful',
                'change': change,
                'total_amount': total_amount,
                'order_summary': f"Sale of {len(cart)} items for ${total_amount:.2f}"
            })

        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Error: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'Invalid request method.'})


def user_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']

     
        user = authenticate(username=username, password=password)

        if user is not None:
            login(request, user)  

           
            if hasattr(user, 'userprofile'):
               
                print(f"Account type for {user.username}: {user.userprofile.account_type}") 
                
               
                if user.userprofile.account_type == 'Admin':
                    return redirect('admin_dashboard')  
                else:
                    return redirect('cashier_pos') 
            else:
               
                UserProfile.objects.create(user=user, account_type='Cashier')
                return redirect('cashier_pos')  

        else:
           
            return render(request, 'login.html', {'error': 'Invalid credentials'})

    return render(request, 'login.html')  

def homepage(request):
    return redirect('user_login')

@login_required
def inventory_management(request):
    items = Item.objects.all()
    total_stock = sum(item.quantity for item in items)  # Sum of all item quantities
    return render(request, 'inventory_management.html', {'items': items, 'total_stock': total_stock})

@login_required
def admin_dashboard(request):
    total_items = Item.objects.count()
    total_sales = SaleItem.objects.aggregate(total=Sum('subtotal'))['total'] or 0
    total_categories = Category.objects.count()
    total_users = User.objects.count()

    
    top_selling_items = (
        SaleItem.objects
        .select_related('item')  
        .values('item__name', 'item__image') 
        .annotate(total_quantity=Sum('quantity'))
        .order_by('-total_quantity')[:3]
    )

    context = {
        'total_items': total_items,
        'total_sales': total_sales,
        'total_categories': total_categories,
        'total_users': total_users,
        'top_selling_items': top_selling_items,
    }

    return render(request, 'admin_dashboard.html', context)




def stock_log(request):
    logs = StockTransaction.objects.select_related('item').order_by('-date')
    return render(request, 'stock_log.html', {'logs': logs})



@login_required
def inventory_management(request):
    items = Item.objects.all()
    categories = Category.objects.all()

    # Filters
    category_id = request.GET.get('category')
    search_query = request.GET.get('q')
    date_filter = request.GET.get('date')

    if category_id:
        items = items.filter(category_id=category_id)

    if search_query:
        items = items.filter(name__icontains=search_query)

    if date_filter:
        try:
            parsed_date = parse_date(date_filter)
            if parsed_date:
                items = items.filter(date_added__date=parsed_date)
        except:
            pass

    # Calculate total stock
    total_stock = sum(item.quantity for item in items)

    context = {
        'items': items,
        'categories': categories,
        'total_stock': total_stock,
    }
    return render(request, 'inventory_management.html', context)

@api_view(['GET'])
def inventory_api(request):
    items = Item.objects.all()
    serializer = ItemSerializer(items, many=True, context={'request': request})
    return Response(serializer.data)

@login_required
def bulk_delete_items(request):
    selected_ids = request.POST.getlist('selected_items')
    if selected_ids:
        Item.objects.filter(id__in=selected_ids).delete()
        messages.success(request, f"{len(selected_ids)} items deleted successfully.")
    else:
        messages.warning(request, "No items selected for deletion.")
    return redirect('inventory_management')

from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from datetime import datetime, timedelta
from .models import Sale

@login_required
def sales_report(request):
    # Get query parameters
    period = request.GET.get('period', 'today').lower()
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    today = datetime.now().date()

    # Initialize sales_data
    sales_data = Sale.objects.none()

    if start_date and end_date:
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d").date()
            end = datetime.strptime(end_date, "%Y-%m-%d").date()
            sales_data = Sale.objects.filter(timestamp__date__range=(start, end))
            selected_period = f"{start} to {end}"
        except ValueError:
            selected_period = "Invalid Date Format"
            sales_data = Sale.objects.none()
    else:
        if period == 'today':
            sales_data = Sale.objects.filter(timestamp__date=today).prefetch_related('items__item')
        elif period == 'week':
            start_of_week = today - timedelta(days=today.weekday())
            sales_data = Sale.objects.filter(timestamp__date__gte=start_of_week)
        elif period == 'month':
            sales_data = Sale.objects.filter(timestamp__month=today.month, timestamp__year=today.year)
        elif period == 'year':
            sales_data = Sale.objects.filter(timestamp__year=today.year)
        else:
            sales_data = Sale.objects.all()
        selected_period = period.capitalize()


    # Calculate the total sales
    total_sales = sum(sale.total_price for sale in sales_data)
    


    return render(request, 'sales_report.html', {
        'selected_period': selected_period,
        'sales_data': sales_data,
        'total_sales': total_sales,
      
        'today': today.strftime('%Y-%m-%d'),
    })


def sales_report_api(request):
    period = request.GET.get('period', 'today').lower()
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    today = datetime.now().date()

    # Initialize sales_data
    sales_data = Sale.objects.none()

    if start_date and end_date:
        sales_data = Sale.objects.filter(timestamp__date__range=(start_date, end_date))
    else:
        if period == 'today':
            sales_data = Sale.objects.filter(timestamp__date=today).prefetch_related('items__item')
        elif period == 'week':
            start_of_week = today - timedelta(days=today.weekday())
            sales_data = Sale.objects.filter(timestamp__date__gte=start_of_week)
        elif period == 'month':
            sales_data = Sale.objects.filter(timestamp__month=today.month, timestamp__year=today.year)
        elif period == 'year':
            sales_data = Sale.objects.filter(timestamp__year=today.year)
        else:
            sales_data = Sale.objects.all()

        selected_period = period.capitalize()

    # Calculate the total sales
    total_sales = sum(sale.total_price for sale in sales_data)

    # Prepare the response data
    sales_data_list = []
    for sale in sales_data:
        sales_data_list.append({
            'timestamp': sale.timestamp,
            'total_price': sale.total_price
        })

    return JsonResponse({
        'selected_period': selected_period,
        'sales_data': sales_data_list,
        'total_sales': total_sales,
    })




def add_item(request):
    if request.method == 'POST':
        form = ItemForm(request.POST, request.FILES)
        if form.is_valid():
            item = form.save(commit=False)

            # Set the quantity directly instead of adding
            quantity_input = int(request.POST.get('quantity', 0))
            item.quantity = quantity_input
            item.save()

            # Record stock transaction (optional)
            StockTransaction.objects.create(
                item=item,
                transaction_type='IN',
                quantity=quantity_input,
                unit=item.unit,
                remarks="Initial stock"
            )

            messages.success(request, f"{item.name} was added with quantity {quantity_input}")
            return redirect('inventory_management')
    else:
        form = ItemForm()

    return render(request, 'add_item.html', {'form': form})


def cashier_pos(request):
    return render(request, 'cashier_pos.html')

def category_management(request):
    if request.method == 'POST':
        # Handle Add
        if 'category_name' in request.POST and 'edit_id' not in request.POST and 'delete_id' not in request.POST:
            category_name = request.POST.get('category_name')
            if category_name:
                Category.objects.create(name=category_name)

        # Handle Edit
        elif 'edit_id' in request.POST:
            category_id = request.POST.get('edit_id')
            category_name = request.POST.get('category_name')
            if category_name:
                category = get_object_or_404(Category, id=category_id)
                category.name = category_name
                category.save()

        # Handle Delete
        elif 'delete_id' in request.POST:
            category_id = request.POST.get('delete_id')
            category = get_object_or_404(Category, id=category_id)
            category.delete()

        return redirect('category_management')

    categories = Category.objects.all()
    return render(request, 'category_management.html', {'categories': categories})

def category_list_api(request):
    categories = Category.objects.all().values('id', 'name')
    return JsonResponse(list(categories), safe=False)

@login_required
def user_management(request):
    if request.method == 'POST':
        form = UserForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)

            raw_password = form.cleaned_data.get('password')
            user.set_password(raw_password)

            account_type_raw = form.cleaned_data['account_type']
            account_type = account_type_raw.capitalize()  # Fix the capitalization (e.g., 'Admin', 'Cashier')

            # Set staff/superuser flags if admin
            if account_type == 'Admin':
                user.is_staff = True  # appears in Django Admin
                #user.is_superuser = True  # optional if you want full admin privileges

            user.save()

            # Save to UserProfile
            UserProfile.objects.create(user=user, account_type=account_type)

            # Show proper success message
            if account_type == 'Admin':
                messages.success(request, 'Admin account created successfully!')
            else:
                messages.success(request, 'Cashier account created successfully!')

            return redirect('user_management')
        else:
            messages.error(request, 'Form has errors. Please check the inputs.')
    else:
        form = UserForm()

    # Handle delete user
    delete_id = request.GET.get('delete_id')
    if delete_id:
        try:
            user_to_delete = User.objects.get(id=delete_id)
            username = user_to_delete.username
            user_to_delete.delete()
            messages.success(request, f"User '{username}' deleted successfully.")
            return redirect('user_management')
        except User.DoesNotExist:
            messages.error(request, "User not found.")

    users = User.objects.all().select_related('userprofile')
    return render(request, 'user_management.html', {'form': form, 'users': users})


def update_item(request, item_id):
    item = get_object_or_404(Item, pk=item_id)
    if request.method == 'POST':
        form = ItemForm(request.POST, instance=item)
        if form.is_valid():
            form.save()
            return redirect('inventory_management')
    else:
        form = ItemForm(instance=item)
    
    return render(request, 'update_item.html', {'form': form, 'item': item})

def delete_item(request, id):
    item = get_object_or_404(Item, id=id)
    item.delete()
    return redirect('inventory_management')


def item_list(request):
    query = request.GET.get('q', '')
    items = Item.objects.all()

    if query:
        items = items.filter(
            Q(name__icontains=query) |
            Q(description__icontains=query) |
            Q(category__name__icontains=query)
        )

    context = {
        'items': items,
    }
    return render(request, 'inventory_management.html', context)



def item_autocomplete(request):
    query = request.GET.get('term', '')
    items = Item.objects.filter(name__icontains=query).values('name', 'id')
    return JsonResponse(list(items), safe=False)

@csrf_exempt
def add_to_cart(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            item_name = data.get('item_name')
            quantity = int(data.get('quantity', 1))

            # Get item from database
            try:
                item = Item.objects.get(name__iexact=item_name)
            except Item.DoesNotExist:
                return JsonResponse({'success': False, 'message': 'Item not found.'})

            # Check if quantity requested exceeds available stock
            if item.quantity < quantity:
                return JsonResponse({'success': False, 'message': f"Only {item.quantity} items available in stock."})

            # Initialize cart in session if not already
            cart = request.session.get('cart', [])

            # Check if item already exists in the session cart
            for cart_item in cart:
                if cart_item['id'] == item.id:
                    cart_item['quantity'] += quantity
                    break
            else:
                # Add new item with item_id properly set
                cart.append({
                    'id': item.id,  # Ensure 'id' matches the 'item_id' in the database
                    'name': item.name,
                    'price': float(item.price),  # Make sure it's JSON serializable
                    'quantity': quantity
                })

            # Save updated cart in session
            request.session['cart'] = cart
            request.session.modified = True

            return JsonResponse({'success': True, 'cart': cart})

        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})

    return JsonResponse({'success': False, 'message': 'Invalid request method'})


@csrf_exempt
def remove_from_cart(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            item_id = data.get('item_id')

            cart = request.session.get('cart', [])
            cart = [item for item in cart if item['id'] != item_id]

            request.session['cart'] = cart
            request.session.modified = True

            return JsonResponse({'success': True, 'cart': cart})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    return JsonResponse({'success': False, 'message': 'Invalid request method'})



def live_search_items(request):
    query = request.GET.get('q', '')
    items = Item.objects.filter(name__icontains=query)
    html = render_to_string('item_rows.html', {'items': items})
    return JsonResponse({'html': html})


def search_item(request, item_name):
    if request.method == 'GET':
        items = Item.objects.filter(name__icontains=item_name)[:10]  # limit for performance
        data = []
        for item in items:
            data.append({
                'id': item.id,
                'name': item.name,
                'price': float(item.price),
                'quantity': item.quantity
            })
        return JsonResponse({'items': data})




@csrf_exempt  # Exempting CSRF for testing (ensure to use CSRF properly in production)
def update_inventory(request, item_id):
    if request.method == 'POST':
        try:
            # Assuming you've already imported your models
            item = Item.objects.get(id=item_id)
            data = json.loads(request.body)
            quantity_sold = data['quantity']

            # Update stock
            if item.stock >= quantity_sold:
                item.stock -= quantity_sold
                item.save()

                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False, 'message': 'Not enough stock.'})
        except Item.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Item not found.'})
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'message': 'Invalid JSON in request.'})
    return JsonResponse({'success': False, 'message': 'Invalid request method.'})

def record_sale(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            sale = Sale.objects.create(
                total_amount=data['totalAmount'],
                received_money=data['receivedMoney'],
                change=data['change']
            )

            for item in data['items']:
                SaleItem.objects.create(
                    sale=sale,
                    item=Item.objects.get(id=item['id']),
                    quantity=item['quantity'],
                    price=item['price']
                )

            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    return JsonResponse({'success': False, 'message': 'Invalid request method.'})



def checkout_view(request):
    # Example: Fetching the cart from the session (adjust based on your logic)
    cart = request.session.get('cart', [])
    
    # Calculate total and other necessary info
    total_amount = sum(item['price'] * item['quantity'] for item in cart)

    return render(request, 'checkout.html', {'cart': cart, 'total_amount': total_amount})



@csrf_exempt  # Disable CSRF for testing (use CSRF protection in production)
def clear_cart(request):
    if request.method == 'POST':
        try:
            # Clear the cart by resetting the session cart to an empty list
            request.session['cart'] = []  
            request.session.modified = True  # Mark the session as modified

            return JsonResponse({'success': True, 'message': 'Cart has been cleared successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})

    return JsonResponse({'success': False, 'message': 'Invalid request method'})


def supplier_list_api(request):
    suppliers = Supplier.objects.all().values(
        'id', 'name', 'contact_person', 'phone', 'email', 'address', 'company', 'created_at'
    )
    suppliers_list = list(suppliers)
    return JsonResponse(suppliers_list, safe=False)

def supplier_list(request):
    suppliers = Supplier.objects.all()
    return render(request, 'supplier_list.html', {'suppliers': suppliers})


def supplier_create(request):
    if request.method == 'POST':
        form = SupplierForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('supplier_list')  # Redirect to the supplier list after saving
    else:
        form = SupplierForm()  # Create a blank form for GET requests

    return render(request, 'supplier_form.html', {'form': form, 'form_title': 'Add Supplier'})

def supplier_update(request, pk):
    supplier = get_object_or_404(Supplier, pk=pk)
    form = SupplierForm(request.POST or None, instance=supplier)
    if form.is_valid():
        form.save()
        return redirect('supplier_list')
    return render(request, 'supplier_form.html', {'form': form, 'form_title': 'Edit Supplier'})

def supplier_detail(request, pk):
    supplier = get_object_or_404(Supplier, pk=pk)
    return render(request, 'supplier_detail.html', {'supplier': supplier})

def supplier_confirm_delete(request, pk):
    supplier = get_object_or_404(Supplier, pk=pk)
    if request.method == 'POST':
        supplier.delete()
        return redirect('supplier_list')
    return render(request, 'supplier_confirm_delete.html', {'supplier': supplier})

def stock_card(request, item_id):
    item = get_object_or_404(Item, id=item_id)
    transactions = item.stock_transactions.all().order_by('-date')  # using the correct related_name
    return render(request, 'stock_card.html', {'item': item, 'transactions': transactions})


def search_items(request):
    query = request.GET.get('q', '')
    items = Item.objects.filter(name__icontains=query)[:10]
    results = [{'name': item.name, 'price': float(item.price), 'description': item.description} for item in items]
    return JsonResponse(results, safe=False)



def get_sales_by_period(period):
    today = now().date()

    if period == "today":
        return Sale.objects.filter(date=today), "Today's Sales"

    elif period == "week":
        start_of_week = today - timedelta(days=today.weekday())
        end_of_week = start_of_week + timedelta(days=6)
        return Sale.objects.filter(date__range=(start_of_week, end_of_week)), "This Week's Sales"

    elif period == "month":
        return Sale.objects.filter(date__month=today.month, date__year=today.year), "This Month's Sales"

    elif period == "year":
        return Sale.objects.filter(date__year=today.year), "This Year's Sales"

    return Sale.objects.all(), "All Sales"


def download_sales_report(request):
    # Get period and format from request
    period = request.GET.get('period', 'today').lower()
    file_format = request.GET.get('format', 'docx').lower()

    # Filter sales and get label
    sales, period_display = get_sales_by_period(period)

    if file_format == 'docx':
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="SST_Inventory_{period_display}.docx"'

        doc = Document()
        doc.add_heading("SST Inventory Management", 0)
        doc.add_heading(f"{period_display} Report", level=1)

        total_sales = 0
        doc.add_paragraph(f"Total Sales: ₱{total_sales:.2f}")

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Time'
        hdr_cells[2].text = 'Subtotal'

        for sale in sales:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sale.date)
            row_cells[1].text = sale.timestamp.strftime("%H:%M:%S")
            row_cells[2].text = f"₱{sale.subtotal:.2f}"
            total_sales += sale.subtotal

        doc.add_paragraph(f"\nTotal Sales: ₱{total_sales:.2f}")
        doc.save(response)
        return response


    elif file_format == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="SST_Sales_Report_{period_display}.pdf"'

        p = canvas.Canvas(response, pagesize=letter)
        width, height = letter

        # Logo
        logo_path = os.path.join(settings.BASE_DIR, 'core', 'static', 'images', 'samsantek-image1.png')
        try:
            p.drawImage(logo_path, 50, height - 80, width=60, height=60)
        except:
            p.setFont("Helvetica", 10)
            p.drawString(50, height - 20, "[Logo not found]")

        # Title and Subheading
        p.setFont("Helvetica-Bold", 18)
        p.drawCentredString(width / 2, height - 50, "SST Inventory Management System")

        p.setFont("Helvetica", 12)
        p.drawCentredString(width / 2, height - 70, f"Sales Report for {period_display}")

        # Line separator
        p.setStrokeColor(colors.black)
        p.line(50, height - 90, width - 50, height - 90)

        y = height - 110

        # Table headers
        p.setFont("Helvetica-Bold", 11)
        p.setFillColor(colors.lightblue)
        p.rect(50, y - 5, width - 100, 20, fill=1, stroke=0)
        p.setFillColor(colors.black)
        p.drawString(60, y, "Date")
        p.drawString(180, y, "Time")
        p.drawString(300, y, "Subtotal (₱)")

        y -= 30
        total_sales = 0

        for i, sale in enumerate(sales):
            if y < 100:
                # New page
                p.showPage()
                y = height - 50
                p.setFont("Helvetica-Bold", 11)
                p.drawString(60, y, "Date")
                p.drawString(180, y, "Time")
                p.drawString(300, y, "Subtotal (₱)")
                y -= 20

            # Alternate row colors
            bg_color = colors.whitesmoke if i % 2 == 0 else colors.lightgrey
            p.setFillColor(bg_color)
            p.rect(50, y - 5, width - 100, 20, fill=1, stroke=0)

            # Sale data
            p.setFillColor(colors.black)
            p.setFont("Helvetica", 10)
            p.drawString(60, y, str(sale.date))
            p.drawString(180, y, sale.timestamp.strftime("%H:%M:%S"))
            p.drawRightString(400, y, f"₱{sale.subtotal:.2f}")

            total_sales += sale.subtotal
            y -= 25

        # Totals
        p.setFont("Helvetica-Bold", 11)
        p.drawString(60, y - 10, "Total Sales:")
        p.drawRightString(400, y - 10, f"₱{total_sales:.2f}")

        # Footer
        p.setFont("Helvetica-Oblique", 8)
        p.setFillColor(colors.darkgray)
        p.drawString(50, 30, "Generated by SST Inventory Management System")
        p.drawRightString(width - 50, 30, f"Page 1")

        p.showPage()
        p.save()
        return response


def stock_in(request, item_id):
    # Step 1: Get the item
    item = get_object_or_404(Item, id=item_id)
    
    # Step 2: Get the quantity to add (e.g., from a form)
    if request.method == "POST":
        quantity_to_add = int(request.POST['quantity'])  # You can get this from a form input
        remarks = request.POST.get('remarks', '')  # Optional remarks
        
        # Ensure that quantity is valid
        if quantity_to_add <= 0:
            messages.error(request, "Quantity to add must be greater than 0")
            return redirect('stock_in', item_id=item.id)
        
        # Step 3: Update the item’s stock quantity
        item.update_quantity(quantity_to_add)
        
        # Step 4: Record the stock transaction as "IN"
        StockTransaction.objects.create(
            item=item,
            transaction_type="IN",  # This is a stock-in transaction
            quantity=quantity_to_add,
            remarks=remarks
        )
        
        messages.success(request, f"Successfully added {quantity_to_add} units of {item.name} to the stock.")
        
        return redirect('stock_summary', item_id=item.id)  # Redirect to the stock summary page

    return render(request, 'stock_in.html', {'item': item})

class LoginView(APIView):
    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')

        if not username or not password:
            return Response({'error': 'Username and password are required.'}, status=status.HTTP_400_BAD_REQUEST)

        user = authenticate(username=username, password=password)

        if user is not None:
            token, created = Token.objects.get_or_create(user=user)

            try:
                role = user.userprofile.account_type
            except UserProfile.DoesNotExist:
                return Response({'error': 'User profile not found.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response({
                'token': token.key,
                'role': role
            }, status=status.HTTP_200_OK)

        return Response({'error': 'Invalid credentials'}, status=status.HTTP_400_BAD_REQUEST)


    
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_info(request):
    user = request.user
    try:
        profile = UserProfile.objects.get(user=user)
        role = profile.account_type
    except UserProfile.DoesNotExist:
        return Response({'error': 'User profile not found'}, status=500)

    return Response({
        'username': user.username,
        'role': role
    })

class DashboardStatsView(APIView):
    def get(self, request):
        # Fetch your stats here
        total_items = Item.objects.count()
        total_sales = Sale.objects.aggregate(total=Sum('total_price'))['total']
        total_categories = Category.objects.count()
        total_users = User.objects.count()

        data = {
            'total_items': total_items,
            'total_sales': total_sales,
            'total_categories': total_categories,
            'total_users': total_users,
        }
        
        return Response(data, status=status.HTTP_200_OK)
    
@api_view(['GET'])
def stock_log_api(request):
    logs = StockTransaction.objects.select_related('item').order_by('-date')
    serializer = StockTransactionSerializer(logs, many=True)
    return Response(serializer.data)